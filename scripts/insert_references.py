"""Append IEEE references to main_scopus.docx after the References heading.

The 'references' paragraph style in this IEEE template auto-numbers entries
as [1], [2], ... via Word's numbering. Therefore each entry below is given
*without* a "[N]" prefix — Word will render the numbers automatically.

Usage:
    1. Open paper/main_scopus.docx, manually delete any existing reference
       entries you don't want.
    2. Run:
           mamba run -n electr-forecast python -m scripts.insert_references
       The script appends entries below the (last) "References" heading.
       Pass --replace-existing to wipe all 'references'-styled paragraphs
       under the heading first.

The script edits paper/main_scopus.docx in place. Make a backup if needed.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


REFERENCES = [
    'T. Hong and S. Fan, "Probabilistic electric load forecasting: A tutorial review," International Journal of Forecasting, vol. 32, no. 3, pp. 914-938, 2016, doi: 10.1016/j.ijforecast.2015.11.011.',
    'W. Kong, Z. Y. Dong, Y. Jia, D. J. Hill, Y. Xu, and Y. Zhang, "Short-term residential load forecasting based on LSTM recurrent neural network," IEEE Transactions on Smart Grid, vol. 10, no. 1, pp. 841-851, 2019, doi: 10.1109/TSG.2017.2753802.',
    'E. Mocanu, P. H. Nguyen, M. Gibescu, and W. L. Kling, "Deep learning for estimating building energy consumption," Sustainable Energy, Grids and Networks, vol. 6, pp. 91-99, 2016, doi: 10.1016/j.segan.2016.02.005.',
    'D. L. Marino, K. Amarasinghe, and M. Manic, "Building energy load forecasting using deep neural networks," in Proc. 42nd Annu. Conf. IEEE Industrial Electronics Society (IECON), 2016, pp. 7046-7051, doi: 10.1109/IECON.2016.7793413.',
    "G. E. P. Box, G. M. Jenkins, G. C. Reinsel, and G. M. Ljung, Time Series Analysis: Forecasting and Control, 5th ed. Hoboken, NJ, USA: Wiley, 2015.",
    "R. J. Hyndman and G. Athanasopoulos, Forecasting: Principles and Practice, 3rd ed. Melbourne, Australia: OTexts, 2021. [Online]. Available: https://otexts.com/fpp3/",
    'R. J. Hyndman and Y. Khandakar, "Automatic time series forecasting: The forecast package for R," Journal of Statistical Software, vol. 27, no. 3, pp. 1-22, 2008, doi: 10.18637/jss.v027.i03.',
    'F. X. Diebold and R. S. Mariano, "Comparing predictive accuracy," Journal of Business & Economic Statistics, vol. 13, no. 3, pp. 253-263, 1995, doi: 10.1080/07350015.1995.10524599.',
    'T. Chen and C. Guestrin, "XGBoost: A scalable tree boosting system," in Proc. 22nd ACM SIGKDD Int. Conf. Knowledge Discovery and Data Mining, 2016, pp. 785-794, doi: 10.1145/2939672.2939785.',
    'S. Hochreiter and J. Schmidhuber, "Long short-term memory," Neural Computation, vol. 9, no. 8, pp. 1735-1780, 1997, doi: 10.1162/neco.1997.9.8.1735.',
    'A. Vaswani, N. Shazeer, N. Parmar, J. Uszkoreit, L. Jones, A. N. Gomez, L. Kaiser, and I. Polosukhin, "Attention is all you need," in Advances in Neural Information Processing Systems (NeurIPS), 2017.',
    'H. Zhou, S. Zhang, J. Peng, S. Zhang, J. Li, H. Xiong, and W. Zhang, "Informer: Beyond efficient transformer for long sequence time-series forecasting," in Proc. AAAI Conf. Artificial Intelligence, 2021, doi: 10.1609/aaai.v35i12.17325.',
    'H. Wu, J. Xu, J. Wang, and M. Long, "Autoformer: Decomposition transformers with auto-correlation for long-term series forecasting," in Advances in Neural Information Processing Systems (NeurIPS), 2021.',
    'T. Zhou, Z. Ma, Q. Wen, X. Wang, L. Sun, and R. Jin, "FEDformer: Frequency enhanced decomposed transformer for long-term series forecasting," in Proc. Int. Conf. Machine Learning (ICML), 2022.',
    'A. Zeng, M. Chen, L. Zhang, and Q. Xu, "Are transformers effective for time series forecasting?" in Proc. AAAI Conf. Artificial Intelligence, 2023, doi: 10.1609/aaai.v37i9.26317.',
    'Y. Nie, N. H. Nguyen, P. Sinthong, and J. Kalagnanam, "A time series is worth 64 words: Long-term forecasting with transformers," in Proc. Int. Conf. Learning Representations (ICLR), 2023.',
    'H. Wu, T. Hu, Y. Liu, H. Zhou, J. Wang, and M. Long, "TimesNet: Temporal 2D-variation modeling for general time series analysis," in Proc. Int. Conf. Learning Representations (ICLR), 2023.',
    'Y. Liu, T. Hu, H. Zhang, H. Wu, S. Wang, L. Ma, and M. Long, "iTransformer: Inverted transformers are effective for time series forecasting," in Proc. Int. Conf. Learning Representations (ICLR), 2024.',
    'Q. Wen, T. Zhou, C. Zhang, W. Chen, Z. Ma, J. Yan, and L. Sun, "Transformers in time series: A survey," in Proc. Int. Joint Conf. Artificial Intelligence (IJCAI), 2023.',
    'A. Das, W. Kong, R. Sen, and Y. Zhou, "A decoder-only foundation model for time-series forecasting," in Proc. Int. Conf. Machine Learning (ICML), 2024.',
    'A. F. Ansari et al., "Chronos: Learning the language of time series," Transactions on Machine Learning Research, 2024.',
    'M. Goswami, K. Szafer, A. Choudhry, Y. Cai, S. Li, and A. Dubrawski, "MOMENT: A family of open time-series foundation models," in Proc. Int. Conf. Machine Learning (ICML), 2024.',
    'G. Woo, C. Liu, A. Kumar, C. Xiong, S. Savarese, and D. Sahoo, "Unified training of universal time series forecasting transformers," in Proc. Int. Conf. Machine Learning (ICML), 2024.',
    'K. Rasul et al., "Lag-Llama: Towards foundation models for probabilistic time series forecasting," 2024, arXiv:2310.08278.',
    'G. Lai, W.-C. Chang, Y. Yang, and H. Liu, "Modeling long- and short-term temporal patterns with deep neural networks," in Proc. 41st Int. ACM SIGIR Conf. Research & Development in Information Retrieval, 2018, pp. 95-104, doi: 10.1145/3209978.3210006.',
    'K. G. Olivares, C. Challu, F. Garza, M. Mergenthaler-Canseco, and A. Dubrawski, "NeuralForecast: User friendly state-of-the-art neural forecasting models," 2022. [Online]. Available: https://github.com/Nixtla/neuralforecast',
    'F. Garza, M. Mergenthaler-Canseco, K. G. Olivares, and C. Challu, "StatsForecast: Lightning fast forecasting with statistical and econometric models," 2022. [Online]. Available: https://github.com/Nixtla/statsforecast',
    'S. Seabold and J. Perktold, "statsmodels: Econometric and statistical modeling with Python," in Proc. 9th Python in Science Conf. (SciPy), 2010, pp. 92-96.',
]


ROOT = Path(__file__).resolve().parent.parent
DEFAULT_DOCX = ROOT / "paper" / "main_scopus.docx"


def find_references_heading(doc):
    """Return the *last* paragraph that is a 'References' heading, or None."""
    target = None
    for p in doc.paragraphs:
        if p.text.strip().lower() == "references" and p.style.name in {
            "Heading 5", "Heading 1", "heading 5", "heading 1",
        }:
            target = p
    return target


def remove_existing_entries(doc, heading) -> int:
    """Delete every paragraph after `heading` that uses the 'references' style.

    Stops at the first non-references paragraph after the heading so it never
    eats unrelated content. Returns the number of paragraphs removed.
    """
    body = heading._p.getparent()
    children = list(body)
    try:
        idx = children.index(heading._p)
    except ValueError:
        return 0
    removed = 0
    i = idx + 1
    while i < len(children):
        el = children[i]
        if el.tag != qn("w:p"):
            break
        pStyle = el.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        if pStyle is None or pStyle.get(qn("w:val")) != "references":
            # Skip empty trailing paragraphs but stop at real content.
            text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
            if text:
                break
            i += 1
            continue
        body.remove(el)
        removed += 1
        children = list(body)
    return removed


def insert_references(doc, heading, entries) -> int:
    """Append each entry as a 'references'-styled paragraph after the heading."""
    style = doc.styles["references"]
    body = heading._p.getparent()
    # Find insertion anchor: place new paragraphs right after the heading.
    anchor = heading._p
    inserted = 0
    parent = doc.paragraphs[0]._parent
    for text in entries:
        new_p = parent.add_paragraph(text, style=style)
        anchor.addnext(new_p._p)
        anchor = new_p._p
        inserted += 1
    return inserted


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX,
                        help="Target docx file (edited in place).")
    parser.add_argument("--replace-existing", action="store_true",
                        help="Delete current 'references'-styled paragraphs "
                             "under the References heading before inserting.")
    args = parser.parse_args()

    if not args.docx.exists():
        raise SystemExit(f"File not found: {args.docx}")

    doc = Document(args.docx)
    heading = find_references_heading(doc)
    if heading is None:
        raise SystemExit("No 'References' heading found. Add a Heading 5 "
                         "(or Heading 1) paragraph with the text 'References' "
                         "and rerun.")

    removed = 0
    if args.replace_existing:
        removed = remove_existing_entries(doc, heading)

    inserted = insert_references(doc, heading, REFERENCES)
    doc.save(args.docx)
    print(f"Heading found      : '{heading.text}' (style {heading.style.name})")
    print(f"Existing removed   : {removed}")
    print(f"References inserted: {inserted}")
    print(f"Saved -> {args.docx}")


if __name__ == "__main__":
    main()
