"""Replace numeric tables in paper/main_scopus.docx with v2 results.

Updates 4 tables in place:
  Table[2] 25x6 — ECL accuracy
  Table[3] 25x6 — ETTh1 accuracy
  Table[4] 49x6 — Computational efficiency
  Table[5] 7x5  — Diebold-Mariano tests

The narrative tables (Table[0] best-model overview, Table[1] family
profile) are touched separately because their content is interpretive.

Cell text is rewritten while preserving paragraph styles and runs, so
the IEEE template's table styling stays intact.
"""

from __future__ import annotations

import argparse
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
RESULTS = ROOT / "results" / "main_results_final.csv"
STAT_TESTS = ROOT / "results" / "stat_tests_final.csv"
DOCX = ROOT / "paper" / "main_scopus.docx"

MODEL_LABEL = {
    "seasonal_naive": "SeasonalNaive",
    "sarima": "SARIMA",
    "xgboost": "XGBoost",
    "dlinear": "DLinear",
    "patchtst": "PatchTST",
    "itransformer": "iTransformer",
    "chronos_bolt_small": "Chronos-Bolt",
    "timesfm": "TimesFM",
}

DETERMINISTIC = {"seasonal_naive", "sarima", "chronos_bolt_small", "timesfm"}
HORIZONS = [24, 96, 168]


def fmt_acc(value: float, scale: str) -> str:
    if not np.isfinite(value):
        return ""
    if scale == "big":
        return f"{value:.2f}"
    return f"{value:.2f}"


def aggregate_accuracy(df: pd.DataFrame, dataset: str) -> pd.DataFrame:
    sub = df[df.dataset == dataset]
    rows = []
    scale = "big" if dataset == "ecl" else "small"
    for horizon in HORIZONS:
        block = sub[sub.horizon == horizon]
        for model, part in block.groupby("model"):
            metrics = {}
            for metric in ["MAE", "RMSE", "sMAPE", "wMAPE"]:
                values = part[metric].to_numpy(dtype=float)
                if model in DETERMINISTIC or len(values) == 1:
                    metrics[metric] = fmt_acc(values.mean(), scale)
                else:
                    metrics[metric] = f"{fmt_acc(values.mean(), scale)} ± {fmt_acc(values.std(ddof=1), scale)}"
            rows.append({
                "model": MODEL_LABEL.get(model, model),
                "model_raw": model,
                "horizon": horizon,
                "MAE_sort": part["MAE"].mean(),
                **metrics,
            })
    out = pd.DataFrame(rows)
    out = out.sort_values(["horizon", "MAE_sort"]).reset_index(drop=True)
    return out


def aggregate_efficiency(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for (dataset, horizon, model), part in df.groupby(["dataset", "horizon", "model"]):
        rows.append({
            "model": MODEL_LABEL.get(model, model),
            "model_raw": model,
            "Data/H": f"{'ECL' if dataset == 'ecl' else 'ETTh1'}/H{horizon}",
            "dataset_order": 0 if dataset == "ecl" else 1,
            "horizon": horizon,
            "Train s": f"{part['train_time_s'].mean():.2f}",
            "Infer ms": f"{part['inference_ms'].mean():.2f}",
            "Params": f"{int(round(part['n_params'].mean())):,}".replace(",", " "),
            "GPU MB": f"{part['peak_gpu_mb'].mean():.1f}",
        })
    out = pd.DataFrame(rows)
    out = out.sort_values(["dataset_order", "horizon", "model"]).reset_index(drop=True)
    return out


def format_pvalue(p: float) -> str:
    if not np.isfinite(p):
        return ""
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"


def aggregate_stat_tests(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    order = [("ecl", 24), ("ecl", 96), ("ecl", 168), ("etth1", 24), ("etth1", 96), ("etth1", 168)]
    for ds, h in order:
        part = df[(df.dataset == ds) & (df.horizon == h)]
        if part.empty:
            continue
        r = part.iloc[0]
        a = MODEL_LABEL.get(r["model_a"], r["model_a"])
        b = MODEL_LABEL.get(r["model_b"], r["model_b"])
        sig = "yes" if bool(r["significant"]) else "no"
        rows.append({
            "Data/H": f"{'ECL' if ds == 'ecl' else 'ETTh1'}/H{h}",
            "Comparison": f"{a} vs {b}",
            "Statistic": f"{float(r['statistic']):.2f}",
            "p-value": format_pvalue(float(r["p_value"])),
            "Sig.": sig,
        })
    return pd.DataFrame(rows)


def replace_cell_text(cell, new_text: str) -> None:
    """Set cell text while preserving the first paragraph's style/runs."""
    if not cell.paragraphs:
        cell.add_paragraph(new_text)
        return
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)
    for extra in cell.paragraphs[1:]:
        extra.text = ""


def write_table(tbl, header: list[str], data_rows: list[list[str]]) -> None:
    expected_rows = 1 + len(data_rows)
    if len(tbl.rows) != expected_rows:
        raise ValueError(
            f"Table row mismatch: expected {expected_rows}, got {len(tbl.rows)}"
        )
    n_cols = len(header)
    for i, cell in enumerate(tbl.rows[0].cells[:n_cols]):
        replace_cell_text(cell, header[i])
    for r, data in enumerate(data_rows, start=1):
        for i, cell in enumerate(tbl.rows[r].cells[:n_cols]):
            replace_cell_text(cell, data[i] if i < len(data) else "")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--results", type=Path, default=RESULTS)
    parser.add_argument("--stat-tests", type=Path, default=STAT_TESTS)
    parser.add_argument("--docx", type=Path, default=DOCX)
    args = parser.parse_args()

    results = pd.read_csv(args.results)
    stat_tests = pd.read_csv(args.stat_tests)

    ecl_acc = aggregate_accuracy(results, "ecl")
    etth1_acc = aggregate_accuracy(results, "etth1")
    efficiency = aggregate_efficiency(results)
    stat_rows = aggregate_stat_tests(stat_tests)

    doc = Document(args.docx)

    acc_header = ["Model", "H", "MAE", "RMSE", "sMAPE (%)", "wMAPE (%)"]
    eff_header = ["Model", "Data/H", "Train s", "Infer ms", "Params", "GPU MB"]
    dm_header = ["Data/H", "Comparison", "Statistic", "p-value", "Sig."]

    ecl_rows = [[r["model"], str(int(r["horizon"])), r["MAE"], r["RMSE"], r["sMAPE"], r["wMAPE"]]
                for _, r in ecl_acc.iterrows()]
    etth1_rows = [[r["model"], str(int(r["horizon"])), r["MAE"], r["RMSE"], r["sMAPE"], r["wMAPE"]]
                  for _, r in etth1_acc.iterrows()]
    eff_rows = [[r["model"], r["Data/H"], r["Train s"], r["Infer ms"], r["Params"], r["GPU MB"]]
                for _, r in efficiency.iterrows()]
    dm_rows = [[r["Data/H"], r["Comparison"], r["Statistic"], r["p-value"], r["Sig."]]
               for _, r in stat_rows.iterrows()]

    summary_header = ["Dataset/H", "Best model", "Close competitors", "Interpretation"]
    summary_rows = [
        ["ECL/H24", "iTransformer", "DLinear, PatchTST", "Supervised NF dominate"],
        ["ECL/H96", "iTransformer", "DLinear, PatchTST", "Trained models far ahead of foundation"],
        ["ECL/H168", "iTransformer", "DLinear, PatchTST", "Trained models far ahead of foundation"],
        ["ETTh1/H24", "Chronos-Bolt", "DLinear, PatchTST", "Zero-shot foundation strongest"],
        ["ETTh1/H96", "DLinear", "Chronos-Bolt, TimesFM", "Linear neural model wins"],
        ["ETTh1/H168", "DLinear", "Chronos-Bolt, TimesFM", "Linear neural model wins"],
    ]

    family_header = ["Family", "Training cost", "Inference profile", "Practical role"]
    family_rows = [
        ["SeasonalNaive/SARIMA", "Low to moderate", "Very low latency", "Reference and fallback"],
        ["XGBoost", "Low", "Low latency", "Supervised tabular reference"],
        ["DLinear/PatchTST/iTransformer", "Moderate (5k steps)", "GPU-friendly", "Trained NF benchmark"],
        ["Chronos/TimesFM", "No task-specific training", "Model-dependent latency", "Cold-start forecasting"],
    ]

    write_table(doc.tables[0], summary_header, summary_rows)
    write_table(doc.tables[1], family_header, family_rows)
    write_table(doc.tables[2], acc_header, ecl_rows)
    write_table(doc.tables[3], acc_header, etth1_rows)
    write_table(doc.tables[4], eff_header, eff_rows)
    write_table(doc.tables[5], dm_header, dm_rows)

    doc.save(args.docx)
    print(f"Wrote {len(ecl_rows)} rows to ECL acc, {len(etth1_rows)} to ETTh1 acc, "
          f"{len(eff_rows)} to efficiency, {len(dm_rows)} to DM tests")
    print(f"Saved -> {args.docx}")


if __name__ == "__main__":
    main()
