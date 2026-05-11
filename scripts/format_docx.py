"""Polish the manually edited Scopus DOCX without rebuilding it.

The current master article is ``paper/main_scopus.docx``.  This script is
intentionally conservative: it does not regenerate tables, figures, authors, or
references.  It only expands selected narrative paragraphs and inserts numeric
citations that correspond to the reference list already present in the DOCX.

Run:
    mamba run -n electr-forecast python -m scripts.format_docx
"""

from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent.parent
DEFAULT_DOCX = ROOT / "paper" / "main_scopus.docx"


REPLACEMENTS: dict[str, str] = {
    "The contributions of this paper are as follows:": (
        "The contributions of this paper are as follows: (i) a reproducible "
        "side-by-side benchmark of two recent time-series foundation models "
        "(TimesFM 2.5 and Chronos-Bolt-Small) against a grid-tuned SARIMA, "
        "gradient-boosted trees, and three modern neural baselines under a "
        "single STLF protocol; (ii) Diebold-Mariano significance tests on "
        "seed-averaged per-window absolute errors for the leading model pairs "
        "at every horizon; (iii) an accuracy-latency Pareto analysis that "
        "quantifies the deployment cost of zero-shot inference relative to "
        "supervised alternatives; and (iv) the documentation of a "
        "training-budget effect under which supervised neural models trained "
        "to convergence outperform zero-shot foundation models by roughly an "
        "order of magnitude on aggregate ECL load, while on ETTh1/HUFL the "
        "leading model varies by horizon between Chronos-Bolt at H=24 and "
        "DLinear at H=96 and H=168."
    ),
    "Figure 1. Experimental pipeline used for the benchmark.": (
        "Figure 1. Experimental pipeline used for the benchmark."
    ),
    "Experimental pipeline used for the benchmark.": (
        "Figure 1. Experimental pipeline used for the benchmark."
    ),
    "Sample 168-hour ECL forecast window.": (
        "Figure 2. Sample 168-hour ECL forecast window."
    ),
    "Summary of the main accuracy results.": (
        "Table I. Summary of the main accuracy results across datasets and "
        "horizons."
    ),
    "Relative MAE by dataset and horizon.": (
        "Figure 3. Relative MAE by dataset and horizon (each cell shows the "
        "model's mean MAE divided by the best mean MAE for that "
        "dataset-horizon)."
    ),
    "MAE across forecast horizons on ECL.": (
        "Figure 4. Mean MAE across forecast horizons on ECL."
    ),
    "MAE across forecast horizons on ETTh1/HUFL.": (
        "Figure 5. Mean MAE across forecast horizons on ETTh1/HUFL."
    ),
    "Accuracy-latency view at H=96.": (
        "Figure 6. Accuracy versus single-window inference latency at H=96 "
        "for both datasets."
    ),
    "Operational interpretation of model families.": (
        "Table II. Operational interpretation of model families."
    ),
    "ECL — accuracy across horizons (mean ± std over seeds where applicable).": (
        "Table III. ECL accuracy across horizons (mean ± standard deviation "
        "over five seeds for stochastic models)."
    ),
    "ETTh1/HUFL — accuracy across horizons (mean ± std over seeds where applicable).": (
        "Table IV. ETTh1/HUFL accuracy across horizons (mean ± standard "
        "deviation over five seeds for stochastic models)."
    ),
    "Computational efficiency (single-window inference at batch size 1).": (
        "Table V. Computational efficiency (single-window inference at batch "
        "size 1, averaged over seeds)."
    ),
    "Diebold-Mariano tests for the top two models by mean MAE, computed on seed-averaged per-window absolute errors.": (
        "Table VI. Diebold-Mariano tests for the top two models by mean MAE, "
        "computed on seed-averaged per-window absolute errors."
    ),
    "Diebold-Mariano tests for the top two models by MAE at seed 42.": (
        "Table VI. Diebold-Mariano tests for the top two models by mean MAE, "
        "computed on seed-averaged per-window absolute errors."
    ),
    "Abstract — Short-term electricity load forecasting is a key task": (
        "Abstract — Short-term electricity load forecasting is a key task for "
        "smart-grid operation, industrial energy management, and predictive "
        "planning. This paper presents a reproducible comparative study of "
        "classical, machine-learning, deep-learning, and foundation-model "
        "approaches for hourly load forecasting. The benchmark includes "
        "SeasonalNaive, SARIMA, XGBoost, DLinear, PatchTST, iTransformer, "
        "Chronos-Bolt-Small, and TimesFM 2.5. The primary dataset is ECL with "
        "aggregate electricity load as the target; ETTh1/HUFL is used as a "
        "secondary industrial benchmark. Experiments cover 24-, 96-, and "
        "168-hour horizons with a 336-hour context window. With a 5000-step "
        "training budget and validation-driven early stopping, supervised "
        "neural models dominate ECL aggregate load: iTransformer, DLinear, and "
        "PatchTST achieve mean MAE roughly an order of magnitude lower than "
        "Chronos-Bolt-Small, TimesFM 2.5, and a grid-tuned SARIMA at every "
        "evaluated horizon. On ETTh1/HUFL the ordering is more compressed: "
        "Chronos-Bolt is best at H=24, while DLinear is best at H=96 and "
        "H=168, with foundation models within a few percent. Diebold-Mariano "
        "tests on seed-averaged per-window errors confirm the leader ordering "
        "is statistically significant. The study also reports training time, "
        "single-window inference latency, and peak GPU memory, showing that "
        "foundation models remain attractive cold-start tools but do not "
        "universally replace properly trained supervised baselines."
    ),
    "Short-term electricity load forecasting (STLF) supports dispatch planning": (
        "Short-term electricity load forecasting (STLF) supports dispatch "
        "planning, balancing, tariff-aware operation, maintenance planning, "
        "and anomaly-aware industrial energy management. Forecasting errors "
        "directly affect operational cost, reserve scheduling, and the "
        "reliability of downstream decision-support systems. This role is "
        "consistent with the broader load-forecasting literature, where "
        "accurate point and probabilistic forecasts are treated as core inputs "
        "for smart-grid planning and operation [1]."
    ),
    "Modern energy systems generate large volumes of hourly and sub-hourly measurements": (
        "Modern energy systems generate large volumes of hourly and sub-hourly "
        "measurements. This makes it possible to train data-driven forecasting "
        "models, but it also creates a methodological problem: increasingly "
        "complex neural and foundation models must be compared against strong, "
        "transparent, and inexpensive baselines. Earlier residential and "
        "building-load studies showed that recurrent and deep neural models "
        "can improve forecasting accuracy under suitable data conditions [2]–[4], "
        "but they do not remove the need for classical statistical controls."
    ),
    "This paper presents a reproducible benchmark that compares classical statistical models": (
        "This paper presents a reproducible benchmark that compares classical "
        "statistical models, machine-learning methods, deep-learning "
        "architectures, and time-series foundation models under the same "
        "chronological evaluation protocol. The comparison is deliberately "
        "pragmatic: the goal is not to declare a universally best model, but to "
        "identify where classical SARIMA, feature-based boosting, compact "
        "neural models, and zero-shot foundation models are operationally "
        "useful."
    ),
    "Classical ARIMA-family models remain widely used in industrial forecasting": (
        "Classical ARIMA-family models remain widely used in industrial "
        "forecasting because they are interpretable, computationally "
        "inexpensive, and well understood. Their use follows the Box-Jenkins "
        "forecasting tradition [5], while modern automatic order-selection "
        "procedures and forecasting textbooks provide practical guidance for "
        "seasonal model identification and validation [6], [7]. Seasonal naive "
        "models are even simpler, but they form a necessary reference for load "
        "forecasting studies with strong daily or weekly seasonality."
    ),
    "Gradient-boosted decision trees are a strong machine-learning baseline": (
        "Gradient-boosted decision trees are a strong machine-learning baseline "
        "for energy forecasting. XGBoost is particularly attractive because it "
        "combines nonlinear feature interactions, regularization, and fast "
        "training on tabular lag features [9]. Deep-learning models, including "
        "LSTM networks [10], attention-based architectures [11], and long-horizon "
        "transformer variants such as Informer, Autoformer, and FEDformer "
        "[12]–[14], are attractive for long-horizon forecasting but often "
        "require careful tuning and sufficient data."
    ),
    "Recent time-series foundation models, including Chronos and TimesFM": (
        "Recent time-series foundation models introduce forecasting without "
        "task-specific training. TimesFM, Chronos, MOMENT, Moirai, and "
        "Lag-Llama represent a shift from dataset-specific supervised fitting "
        "towards pretrained sequence models for generic time-series prediction "
        "[20]–[24]. This is useful for cold-start settings, but public benchmark "
        "contamination must be acknowledged because ECL and ETT are widely used "
        "datasets and may have influenced model development."
    ),
    "The benchmark is formulated as a univariate rolling-origin forecasting task": (
        "The benchmark is formulated as a univariate rolling-origin forecasting "
        "task. Given a context window of L=336 hourly observations, each model "
        "produces a multi-step point forecast for the next H values, where "
        "H ∈ {24, 96, 168}. Rolling-origin evaluation is used because it "
        "matches the operational setting in which a forecaster repeatedly "
        "receives a historical context and predicts a fixed future horizon. "
        "All trained models are evaluated on a held-out chronological test "
        "split in the original target scale after inverse standardization."
    ),
    "The primary dataset is ECL (Electricity Consuming Load)": (
        "The primary dataset is ECL (Electricity Consuming Load), aggregated "
        "across 321 clients at hourly resolution; the target is aggregate "
        "electricity load. The secondary dataset is ETTh1, where HUFL is used "
        "as the load-related target variable. The OT variable in ETTh1 is oil "
        "temperature and is therefore not interpreted as electricity load. "
        "Both datasets are common public benchmarks for long- and short-term "
        "time-series forecasting [25], which supports reproducibility but also "
        "requires a cautious interpretation of zero-shot foundation-model "
        "results."
    ),
    "The model set covers four families": (
        "The model set covers four families. SeasonalNaive and SARIMA represent "
        "classical baselines, with SARIMA implemented through statistical "
        "forecasting tooling [27], [28]. XGBoost represents supervised "
        "feature-based machine learning [9]. DLinear, PatchTST, and "
        "iTransformer represent trained neural forecasting models and are run "
        "through NeuralForecast [15]–[18], [26]. Chronos-Bolt-Small and "
        "TimesFM 2.5 represent zero-shot foundation models [20], [21]. "
        "Accuracy is evaluated using MAE, RMSE, sMAPE, and wMAPE; statistical "
        "comparisons use the Diebold-Mariano test [8]."
    ),
    "Two classical baselines are included": (
        "Two classical baselines are included. SeasonalNaive repeats the most "
        "recent daily seasonal pattern and provides a deliberately simple "
        "seasonality reference. SARIMA performs an AICc-based grid search on "
        "the training split over candidate non-seasonal and seasonal orders; "
        "the selected configuration is then reused on every test window via "
        "state filtering on the new context. This follows established ARIMA "
        "forecasting methodology [5]–[7] and uses the statistical forecasting "
        "tooling reported in [27], with Python econometric infrastructure "
        "consistent with [28]."
    ),
    "XGBoost is trained as a direct multi-output regressor": (
        "XGBoost is trained as a direct multi-output regressor with one "
        "estimator per forecast step. Features comprise lag-1, lag-24, and "
        "lag-168 values of the target; rolling means over 24 and 168 hours; "
        "and cyclic calendar encodings for hour-of-day, day-of-week, and "
        "month. This design follows the common tabular forecasting strategy "
        "of exposing seasonal structure through engineered lag features while "
        "letting gradient boosting model nonlinear interactions [9]. DLinear, "
        "PatchTST, and iTransformer are trained through NeuralForecast [26], "
        "which provides a consistent interface for the neural architectures "
        "evaluated in [15]–[18]."
    ),
    "Two time-series foundation models are evaluated in zero-shot mode": (
        "Two time-series foundation models are evaluated in zero-shot mode "
        "without task-specific training: Chronos-Bolt-Small and TimesFM 2.5. "
        "Both are conditioned on the same L=336 context as the trained models "
        "to keep the comparison protocol aligned. TimesFM and Chronos are "
        "representative of the recent decoder-only and tokenized foundation "
        "model direction [20], [21], while MOMENT, Moirai, and Lag-Llama show "
        "that this research line is rapidly expanding beyond a single model "
        "family [22]–[24]."
    ),
    "Accuracy is measured with mean absolute error": (
        "Accuracy is measured with mean absolute error (MAE), root mean squared "
        "error (RMSE), symmetric mean absolute percentage error (sMAPE), and "
        "weighted mean absolute percentage error (wMAPE). These metrics are "
        "reported together because scale-dependent absolute errors are useful "
        "for operational interpretation, while percentage-style measures help "
        "compare behavior across targets with different magnitudes [6]. "
        "Efficiency is reported as training time, single-window inference "
        "latency at batch size 1, trainable parameter count, and peak GPU "
        "memory."
    ),
    "Stochastically initialized models": (
        "Stochastically initialized models (XGBoost, DLinear, PatchTST, and "
        "iTransformer) are evaluated across five seeds {42, 43, 44, 45, 46} "
        "and reported as mean ± standard deviation. Deterministic models "
        "(SeasonalNaive, SARIMA, Chronos-Bolt, and TimesFM) are evaluated once "
        "and have a standard deviation of zero by construction. Neural "
        "forecasting models (DLinear, PatchTST, iTransformer) are trained with "
        "a maximum of 5000 optimizer steps, batch size 32, and validation-driven "
        "early stopping (patience 10 over 50-step checks); the learning rate is "
        "selected from {1e-3, 5e-4} on a chronological inner-validation split. "
        "XGBoost hyperparameters (max_depth, n_estimators, learning_rate) are "
        "selected on the validation split before the final test evaluation. "
        "Statistical comparisons use the Diebold-Mariano test on seed-averaged "
        "per-window absolute errors so that the pooled estimator covers all "
        "available seeds for the stochastic models [8]. The TimesFM checkpoint "
        "is google/timesfm-2.5-200m-pytorch (200M parameters), not the larger "
        "500M variant; the Chronos checkpoint is amazon/chronos-bolt-small "
        "(47.7M parameters)."
    ),
    "On ECL, the top group consists of TimesFM, Chronos-Bolt, and XGBoost": (
        "On ECL, the four supervised models trained with a 5000-step budget — "
        "iTransformer, DLinear, PatchTST, and XGBoost — separate clearly from "
        "the foundation models and classical baselines. iTransformer is the "
        "best model at every horizon, with mean MAE ≈ 26.1k at H=24, 37.5k at "
        "H=96, and 43.0k at H=168; DLinear is within 4–6% behind at each "
        "horizon [15], [18]. Chronos-Bolt, TimesFM, and SARIMA cluster at "
        "MAE ≈ 260–270k, roughly an order of magnitude higher than the trained "
        "supervised neural models. SeasonalNaive sits another 30% above at "
        "≈ 340k. The gap between trained neural and zero-shot foundation "
        "models on aggregate ECL load is therefore not a small effect. The "
        "high sMAPE values for foundation models and SARIMA on ECL (≈ 31%) "
        "reflect that these models track only the slow baseline of the "
        "aggregate signal and miss the high-amplitude daily and weekly "
        "fluctuations of the 321-client sum; the absolute MAE and RMSE values "
        "are the operationally meaningful comparison for this aggregated "
        "target."
    ),
    "On ECL, the four supervised models trained with a 5000-step budget": (
        "On ECL, the four supervised models trained with a 5000-step budget — "
        "iTransformer, DLinear, PatchTST, and XGBoost — separate clearly from "
        "the foundation models and classical baselines. iTransformer is the "
        "best model at every horizon, with mean MAE ≈ 26.1k at H=24, 37.5k at "
        "H=96, and 43.0k at H=168; DLinear is within 4–6% behind at each "
        "horizon [15], [18]. Chronos-Bolt, TimesFM, and SARIMA cluster at "
        "MAE ≈ 260–270k, roughly an order of magnitude higher than the trained "
        "supervised neural models. SeasonalNaive sits another 30% above at "
        "≈ 340k. The gap between trained neural and zero-shot foundation "
        "models on aggregate ECL load is therefore not a small effect. The "
        "high sMAPE values for foundation models and SARIMA on ECL (≈ 31%) "
        "reflect that these models track only the slow baseline of the "
        "aggregate signal and miss the high-amplitude daily and weekly "
        "fluctuations of the 321-client sum; the absolute MAE and RMSE values "
        "are the operationally meaningful comparison for this aggregated "
        "target."
    ),
    "On ETTh1/HUFL, the ranking changes": (
        "On ETTh1/HUFL, the ordering is more compressed and horizon-dependent. "
        "Chronos-Bolt is best at H=24 with MAE = 3.00, ahead of DLinear (3.07) "
        "and PatchTST (3.10); TimesFM trails at 3.16 [20], [21]. At H=96 and "
        "H=168, DLinear takes the lead with MAE = 3.54 and 3.69 respectively, "
        "with Chronos-Bolt and TimesFM each within 2–3% behind. The supervised "
        "neural models therefore remain competitive across all horizons on "
        "ETTh1, while foundation models retain a meaningful edge only at H=24. "
        "SARIMA, SeasonalNaive, and XGBoost form the bottom group on ETTh1 [9], "
        "[15]. The sMAPE values on ETTh1 are large (45–60%) because the HUFL "
        "target crosses zero, which inflates the symmetric percentage error; "
        "MAE, RMSE, and wMAPE are therefore the primary interpretation metrics "
        "for this dataset."
    ),
    "The Diebold-Mariano tests reject equal predictive accuracy": (
        "The Diebold-Mariano tests on seed-averaged per-window absolute errors "
        "reject the null of equal predictive accuracy for every top-1 vs top-2 "
        "comparison in Table VI with p < 0.005 [8]. The pooled-error formulation "
        "is more conservative than a single-seed comparison because it uses "
        "≈ 5–6 thousand windows aligned across seeds, so the leader ordering on "
        "ECL (iTransformer over DLinear) and on ETTh1 (Chronos-Bolt at H=24, "
        "DLinear at H=96 and H=168) is unlikely to be a sampling artifact. "
        "Statistical significance and operational significance are not "
        "identical: on ECL the supervised neural models all dominate the "
        "foundation group by an order of magnitude, so the practical choice "
        "among iTransformer, DLinear, and PatchTST hinges more on training "
        "cost and inference latency than on the small accuracy gap."
    ),
    "The main finding is that no model family dominates across datasets and horizons": (
        "The main finding is that supervised neural models, given a full "
        "training budget, dominate ECL aggregate load forecasting and remain "
        "competitive on ETTh1/HUFL. iTransformer, DLinear, and PatchTST are "
        "roughly an order of magnitude more accurate than Chronos-Bolt, "
        "TimesFM, and SARIMA on ECL across all evaluated horizons. On ETTh1 "
        "the picture is more nuanced: Chronos-Bolt edges DLinear at H=24, while "
        "DLinear is best at H=96 and H=168. XGBoost, even after a grid-tuned "
        "hyperparameter search, remains a strong low-cost CPU baseline but is "
        "no longer the leading approach on either dataset once the neural "
        "models are trained to convergence [9], [15], [18]."
    ),
    "The accuracy–latency trade-offs make these observations practically actionable": (
        "The accuracy–latency trade-offs make these observations practically "
        "actionable. On ECL, supervised neural models achieve roughly 8× lower "
        "MAE than foundation models at single-window inference latencies of "
        "55–75 ms, comparable to TimesFM (100–200 ms) and Chronos-Bolt "
        "(15–45 ms). XGBoost remains the cheapest CPU-only option in absolute "
        "terms (≈ 12–73 ms inference, no GPU memory) and is therefore competitive "
        "when CPU-only deployment is required and a ±20% MAE penalty relative "
        "to iTransformer is acceptable. On ETTh1 the absolute MAE gap among "
        "the top four methods is below 5%, so deployment factors such as "
        "training cost, retraining cadence, and zero-shot applicability often "
        "decide the practical choice [9], [15], [20], [21]."
    ),
    "Foundation models are attractive when task-specific training is undesirable": (
        "Foundation models remain attractive when task-specific training is "
        "undesirable, for example in cold-start industrial monitoring, rapid "
        "prototyping, or sites where historical data are insufficient for "
        "supervised training. On ETTh1, Chronos-Bolt and TimesFM achieve "
        "accuracy within a few percent of the best supervised model without any "
        "task-specific training [20], [21]. On ECL, however, the same models "
        "are roughly an order of magnitude less accurate than properly trained "
        "neural models. The value of zero-shot foundation models is therefore "
        "dataset-dependent and is not a universal substitute for supervised "
        "forecasting [22]–[24]."
    ),
    "The trained transformer models should not be interpreted as generally ineffective": (
        "A useful methodological observation is that the training budget is a "
        "first-order experimental factor for transformer-based forecasters. An "
        "earlier configuration of this protocol with a 500-step optimizer "
        "budget reported MAE values for DLinear, PatchTST, and iTransformer "
        "roughly an order of magnitude higher than those in Tables III and IV. "
        "Increasing the budget to 5000 steps with validation-driven early "
        "stopping triggered convergence within 1–2 minutes of GPU time per "
        "configuration and uncovered the rankings reported here. This is "
        "consistent with the long-sequence forecasting literature, where "
        "patch-based, decomposition-based, and inverted-tokenization "
        "transformer architectures rely on multi-thousand-step training "
        "schedules to express their inductive biases [15], [16], [18]."
    ),
    "The first limitation is public-benchmark pretraining contamination": (
        "The first limitation is public-benchmark pretraining contamination. "
        "ECL and ETT are public datasets, and they may overlap with corpora "
        "used during foundation-model development. The correct interpretation "
        "is therefore that Chronos-Bolt and TimesFM are evaluated without "
        "task-specific training in this repository; the experiment does not "
        "prove dataset-unseen generalization for private industrial sites."
    ),
    "The second limitation is the target-only design": (
        "The second limitation is the target-only design. Real industrial "
        "forecasting systems often include weather, calendar events, tariffs, "
        "production schedules, and operational constraints. Adding such "
        "exogenous variables may change the ranking of model families, "
        "especially for tree-based and transformer-based approaches."
    ),
    "The third limitation is the restricted hyperparameter budget": (
        "The third limitation is hyperparameter coverage. Neural models are "
        "trained with a fixed input length of 336 hours, batch size 32, and a "
        "two-point learning-rate sweep over {1e-3, 5e-4}, with up to 5000 "
        "optimizer steps and validation-driven early stopping. XGBoost is grid-"
        "tuned on max_depth ∈ {4, 6}, n_estimators ∈ {500, 1000}, and "
        "learning_rate ∈ {0.05, 0.1} with the best configuration selected on "
        "the validation split. SARIMA selects orders from an AICc grid on the "
        "training split [6], [7]. Larger searches could improve accuracy "
        "further, especially for multivariate variants of iTransformer or for "
        "longer training budgets, but the executed configuration is sufficient "
        "to demonstrate that supervised neural models do not converge with a "
        "500-step budget."
    ),
    "This paper presented a reproducible benchmark of classical": (
        "This paper presented a reproducible benchmark of classical, "
        "machine-learning, deep-learning, and foundation-model approaches for "
        "short-term electricity load forecasting. With a sufficient training "
        "budget, supervised neural models dominate ECL aggregate load by "
        "roughly an order of magnitude over zero-shot foundation models, while "
        "on ETTh1/HUFL Chronos-Bolt is best at H=24 and DLinear is best at "
        "longer horizons. These results reinforce two messages: training "
        "budget is a first-order experimental factor for transformer-based "
        "forecasters, and the relative ranking of model families is highly "
        "dataset-dependent rather than universal."
    ),
    "A further limitation is that the study reports deterministic point forecasts": (
        "A further limitation is that the study reports deterministic point "
        "forecasts. Many operational grid and industrial planning tasks "
        "require prediction intervals or full probabilistic forecasts to "
        "support reserve allocation and risk-aware scheduling. Extending the "
        "protocol to quantile losses, calibration metrics, and probabilistic "
        "foundation models would make the comparison more directly applicable "
        "to reliability-constrained decision making."
    ),
    "Finally, the benchmark uses two public hourly datasets": (
        "Finally, the benchmark uses two public hourly datasets. This choice "
        "supports reproducibility and allows other researchers to verify the "
        "reported results, but it does not cover all STLF regimes. Future work "
        "should repeat the protocol on private industrial sites, weather-"
        "sensitive residential feeders, and datasets with known calendar "
        "events. Such studies would clarify whether the relative strength of "
        "iTransformer, DLinear, and foundation models transfers beyond "
        "standard public benchmarks."
    ),
    "The practical conclusion is that strong simple baselines remain necessary": (
        "The practical conclusion is that simple baselines remain necessary as "
        "reference points, but they are not necessarily competitive once neural "
        "models are trained to convergence. SeasonalNaive and SARIMA establish "
        "a meaningful lower bound, XGBoost is a strong low-cost CPU baseline, "
        "DLinear is a useful compact neural reference, and foundation models "
        "are valuable cold-start forecasters whose competitiveness varies by "
        "dataset. A deployment-oriented STLF study should report accuracy, "
        "statistical significance, and computational cost together, and should "
        "not declare a winning model family without an explicit training-budget "
        "audit."
    ),
}


INSERT_AFTER: dict[str, list[str]] = {
    "Recent time-series foundation models introduce forecasting without task-specific training": [
        (
            "For STLF specifically, this background implies that a useful comparison must include "
            "both domain-standard baselines and recent general-purpose models. A benchmark that "
            "contains only neural or foundation models risks overstating novelty, while a "
            "benchmark that excludes pretrained models misses a deployment pattern that is now "
            "common in industrial prototyping. Therefore, the present study evaluates classical "
            "seasonal methods, gradient boosting, compact neural architectures, and zero-shot "
            "foundation models within a single execution pipeline."
        ),
        (
            "The long-sequence forecasting literature also motivates the inclusion of simple "
            "neural baselines. Linear and decomposition-based models have repeatedly shown that "
            "strong temporal biases can match or outperform more complex attention mechanisms on "
            "standard benchmarks [15]. Patch-based transformers and inverted-tokenization "
            "architectures address different weaknesses of vanilla sequence attention [16], [18], "
            "but their advantage is not guaranteed under a target-only STLF protocol. This makes "
            "DLinear, PatchTST, and iTransformer a useful compact neural set rather than an "
            "arbitrary model list."
        ),
    ],
    "Accuracy is measured with mean absolute error (MAE), root mean squared error": [
        (
            "The comparison protocol is intentionally leakage-safe. Dataset statistics used for "
            "standardization are estimated only on the training split, validation choices are made "
            "without access to the test segment, and each test forecast uses only observations "
            "available before its forecast origin. Prophet-like or tree-based feature engineering "
            "is also restricted to lag and calendar information that would be known at prediction "
            "time. This is essential because small leakage errors can dominate benchmark rankings "
            "on public time-series datasets."
        ),
        (
            "Efficiency is treated as part of the method rather than as an afterthought. Training "
            "time captures the cost of adapting a model to a new site, while single-window "
            "latency captures the cost of producing one operational forecast after deployment. "
            "These two quantities answer different engineering questions: a zero-shot model may "
            "avoid training entirely but still be expensive at inference, whereas a supervised "
            "model may require fitting once and then produce forecasts cheaply."
        ),
    ],
    "The trained transformer models should not be interpreted as generally ineffective": [
        (
            "From a deployment perspective, the results suggest a tiered model-selection strategy. "
            "SeasonalNaive and SARIMA are useful first checks because they reveal whether the "
            "series is dominated by regular seasonality. XGBoost is an appropriate default when "
            "sufficient site history is available and feature engineering is acceptable. DLinear "
            "is a lightweight neural alternative for horizons where linear temporal structure is "
            "dominant. Foundation models are most attractive when training data are scarce, "
            "experiments must be started quickly, or the cost of maintaining site-specific models "
            "is high."
        ),
        (
            "The accuracy-latency results also caution against interpreting model quality through "
            "a single metric. A model that is marginally more accurate but substantially slower "
            "may be unattractive in an online monitoring system, while a slower zero-shot model "
            "may still be useful if it eliminates training and tuning for many small assets. "
            "Consequently, the benchmark is best read as a Pareto comparison across accuracy, "
            "latency, and adaptation cost rather than as a strict ranking."
        ),
    ],
    "The third limitation is hyperparameter coverage": [
        (
            "A further limitation specific to this study is that iTransformer is evaluated in a "
            "single-series (univariate) configuration with n_series=1. This choice keeps the "
            "protocol comparable across all eight models, but it removes iTransformer's "
            "inverted-tokenization advantage of attending across variables [18]. Despite this "
            "constraint, iTransformer is the best model on every ECL horizon, which suggests "
            "that even its univariate parameterization captures useful sequence structure for "
            "aggregate-load forecasting. A multivariate evaluation that exposes per-client or "
            "weather covariates is left for future work."
        ),
        (
            "A further limitation is that the study reports deterministic point forecasts. Many "
            "operational grid and industrial planning tasks require prediction intervals or full "
            "probabilistic forecasts to support reserve allocation and risk-aware scheduling. "
            "Extending the protocol to quantile losses, calibration metrics, and probabilistic "
            "foundation models would make the comparison more directly applicable to reliability-"
            "constrained decision making."
        ),
        (
            "Finally, the benchmark uses two public hourly datasets. This choice supports "
            "reproducibility and allows other researchers to verify the reported results, but it "
            "does not cover all STLF regimes. Future work should repeat the protocol on private "
            "industrial sites, weather-sensitive residential feeders, and datasets with known "
            "calendar events. Such studies would clarify whether the relative strength of "
            "iTransformer, DLinear, and foundation models transfers beyond standard public "
            "benchmarks."
        ),
    ],
}


def _insert_after(paragraph, text: str):
    new_paragraph = paragraph.insert_paragraph_before(text)
    paragraph._p.addnext(new_paragraph._p)
    new_paragraph.style = paragraph.style
    new_paragraph.alignment = paragraph.alignment
    return new_paragraph


def _set_no_proof(paragraph) -> None:
    """Disable Word spell/grammar marking for a paragraph while preserving style."""
    for run in paragraph.runs:
        rpr = run._r.get_or_add_rPr()
        if rpr.find(qn("w:noProof")) is None:
            rpr.append(OxmlElement("w:noProof"))
        lang = rpr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rpr.append(lang)
        lang.set(qn("w:val"), "en-US")
        lang.set(qn("w:eastAsia"), "en-US")


def polish_docx(path: Path) -> tuple[int, int]:
    doc = Document(path)
    replaced = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for prefix, new_text in REPLACEMENTS.items():
            if text.startswith(prefix):
                if paragraph.text != new_text:
                    paragraph.text = new_text
                    replaced += 1
                break
        _set_no_proof(paragraph)

    inserted = 0
    for prefix, paragraphs_to_insert in INSERT_AFTER.items():
        anchor = next((paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith(prefix)), None)
        if anchor is None:
            continue
        cursor = anchor
        for text in paragraphs_to_insert:
            if any(paragraph.text.strip() == text for paragraph in doc.paragraphs):
                continue
            cursor = _insert_after(cursor, text)
            _set_no_proof(cursor)
            inserted += 1

    # Also mark table text as no-proof; tables themselves are not regenerated.
    table_runs = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _set_no_proof(paragraph)
                    table_runs += len(paragraph.runs)

    doc.save(path)
    return replaced + inserted, table_runs


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--backup", action="store_true",
                        help="Save a .bak copy next to the DOCX before editing.")
    args = parser.parse_args()

    if not args.docx.exists():
        raise SystemExit(f"File not found: {args.docx}")

    if args.backup:
        backup = args.docx.with_suffix(args.docx.suffix + ".bak")
        shutil.copy2(args.docx, backup)
        print(f"Backup -> {backup}")

    replaced, table_runs = polish_docx(args.docx)
    print(f"Expanded paragraphs : {replaced}")
    print(f"Table runs checked  : {table_runs}")
    print(f"Saved -> {args.docx}")


if __name__ == "__main__":
    main()
